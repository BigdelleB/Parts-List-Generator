import pandas as pd
import docx
from docx import Document

def main(og, new):
	was_is_list = Document()


	#create the data frames
	og_df = pd.read_excel(og, usecols="A:F")
	new_df = pd.read_excel(new, usecols="A:F")

	og_len = len(og_df.index)
	new_len = len(new_df.index)

	#store the differences found here
	was = []
	is_ = []


	end_condition = og_len

	#if the new list is longer
	if new_len>og_len:
		is_.append(new_df.iloc[og_len+1:new_len, :])
		was.append("BLANK")*(new_len-og_len)


	elif og_len>new_len:
		was.append(og_df.iloc[new_len+1:og_len, :])
		is_.append("BLANK")*(og_len-new_len)
		end_condition = new_len


	'''we are going to start with the original list and just compare line by line, at a difference we store the value'''

	k = 0

	while k < end_condition:
		row_og = og_df.loc[k, : ]
		row_new = new_df.loc[k, : ]



		if not row_og.equals(row_new):

			was_is_list.add_paragraph("Item "+str(row_og["ITEM"]) +" was:")


			t_was = was_is_list.add_table(rows = 2, cols = 6)


#confused by the word issues below 
			# add the header rows.
			for j in range(6):
				t_was.cell(0,j).text = str(row_og.index[j])
				t_was.cell(1,j).text = str(row_og.iloc[j])
#confused by the word issues above 

			was_is_list.add_paragraph("Item "+str(row_new["ITEM"]) +" is:")
			t_is = was_is_list.add_table(rows = 2, cols = 6)

			for j in range(6):
				t_is.cell(0,j).text = str(row_new.index[j])
				t_is.cell(1,j).text = str(row_new.iloc[j])


			was.append(row_og)
			is_.append(row_new)


		k = k+1


	was_df = pd.DataFrame(was)
	is_df =  pd.DataFrame(is_)
	
	was_is_list.save('./WAS-IS.docx')




